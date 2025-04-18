"""
Maritime Vessel Safety Systems - Automated Compliance Report Generator

This script reads test result outputs from the automated testing framework and generates 
a formatted compliance report in Microsoft Word (.docx) format. The report includes a summary 
of overall compliance (how many tests passed/failed) and detailed evidence for each test case, 
traced to its requirement.

Inputs:
- Test results in JSON or CSV format. These could be individual JSON files (one per test) in a 
  directory or a single aggregated file. Each record should contain at least:
  test_id, requirement_id (or requirement), result (pass/fail indicator), and details (evidence).
- An optional Word template (.docx) that contains company-specific formatting (logos, headers, etc.).

Output:
- A Word document (.docx) with the compliance report. Users can convert this to PDF as needed.

Usage:
  python compliance_report_generator.py --input <results_dir_or_file> --output <report.docx> [--template <template.docx>]

Example:
  $ python compliance_report_generator.py --input "./test_results" --output "Compliance_Report.docx" --template "Compliance_Template.docx"

Dependencies:
  - python-docx (install via pip) for generating Word documents&#8203;:contentReference[oaicite:7]{index=7}.
"""
import os
import sys
import json
import csv
import logging
import argparse
from datetime import datetime

from docx import Document
from docx.shared import RGBColor

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("compliance_report.log"),
        logging.StreamHandler(sys.stdout)
    ]
)

def load_test_results(input_path):
    """Load test results from a JSON/CSV file or a directory of JSON files. Returns a list of dicts."""
    results = []
    if os.path.isdir(input_path):
        logging.info(f"Loading all JSON results from directory: {input_path}")
        for filename in os.listdir(input_path):
            if filename.lower().endswith(".json"):
                filepath = os.path.join(input_path, filename)
                try:
                    with open(filepath, 'r') as f:
                        data = json.load(f)
                except Exception as e:
                    logging.warning(f"Skipping {filename}: error reading JSON - {e}")
                    continue
                # Each JSON file may contain a single result (dict) or a list of results
                if isinstance(data, dict):
                    results.append(data)
                elif isinstance(data, list):
                    for entry in data:
                        if isinstance(entry, dict):
                            results.append(entry)
    else:
        # Single file input
        if not os.path.exists(input_path):
            logging.error(f"Input file {input_path} not found.")
            sys.exit(1)
        logging.info(f"Loading test results from file: {input_path}")
        if input_path.lower().endswith(".json"):
            with open(input_path, 'r') as f:
                data = json.load(f)
            if isinstance(data, dict):
                results.append(data)
            elif isinstance(data, list):
                results.extend([entry for entry in data if isinstance(entry, dict)])
        elif input_path.lower().endswith(".csv"):
            with open(input_path, newline='') as f:
                reader = csv.DictReader(f)
                for row in reader:
                    # Convert the 'result' field to boolean or standardized text
                    if "result" in row:
                        val = row["result"]
                        if isinstance(val, str):
                            val_low = val.strip().lower()
                            if val_low in ["pass", "passed", "true", "1", "yes", "y"]:
                                row["result"] = True
                            elif val_low in ["fail", "failed", "false", "0", "no", "n"]:
                                row["result"] = False
                    results.append(row)
        else:
            logging.error("Unsupported file format. Please provide .json or .csv or a directory of JSON files.")
            sys.exit(1)
    logging.info(f"Loaded {len(results)} test result records.")
    return results

def generate_report(results, output_path, template_path=None):
    """Generate a compliance report Word document from test results."""
    # Use a template if provided, otherwise create a new document
    if template_path:
        if not os.path.exists(template_path):
            logging.error(f"Template file {template_path} not found.")
            sys.exit(1)
        logging.info(f"Using template document: {template_path}")
        document = Document(template_path)
    else:
        document = Document()  # new blank document
    # Document title and generation timestamp
    document.add_heading("Automated Compliance Test Report", level=0)
    report_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    document.add_paragraph(f"Report generated on: {report_date}")
    # Summary of results
    total_tests = len(results)
    passed_tests = 0
    failed_tests = 0
    failed_requirements = []  # to collect IDs of failed requirements
    requirements_tested = set()
    for res in results:
        req_id = res.get("requirement_id") or res.get("requirement") or ""
        if req_id:
            requirements_tested.add(req_id)
        # Determine if test passed
        res_val = res.get("result")
        if isinstance(res_val, bool):
            passed = res_val
        elif isinstance(res_val, str):
            passed = res_val.strip().lower() in ["pass", "passed", "true", "yes", "y", "1"]
        else:
            passed = bool(res_val)
        if passed:
            passed_tests += 1
        else:
            failed_tests += 1
            if req_id:
                failed_requirements.append(str(req_id))
    summary_para = document.add_paragraph()
    # Compose summary sentence
    summary_text = f"{total_tests} tests were executed, covering {len(requirements_tested)} requirements. "
    if failed_tests == 0:
        summary_text += "All tests passed, indicating full compliance with the tested requirements."
    else:
        summary_text += f"{passed_tests} passed, {failed_tests} failed. The system is **NOT** fully compliant."
    summary_para.add_run(summary_text)
    if failed_tests > 0:
        # List the requirements that failed (non-compliant requirements)
        unique_failed_reqs = sorted(set(failed_requirements))
        if unique_failed_reqs:
            non_compliant_text = "Non-compliant requirements: " + ", ".join(unique_failed_reqs)
            document.add_paragraph(non_compliant_text)
    # Detailed results table
    logging.info("Generating detailed results table in the report.")
    table = document.add_table(rows=1, cols=4)
    # Set table headers
    header_cells = table.rows[0].cells
    header_titles = ["Test ID", "Requirement", "Result", "Details"]
    for idx, title in enumerate(header_titles):
        header_cells[idx].text = title
        # Make header text bold
        for run in header_cells[idx].paragraphs[0].runs:
            run.bold = True
    # Populate table rows with each test result
    for res in results:
        test_id = res.get("test_id", "N/A")
        req_id = res.get("requirement_id") or res.get("requirement") or "N/A"
        res_val = res.get("result")
        if isinstance(res_val, bool):
            passed = res_val
        elif isinstance(res_val, str):
            passed = res_val.strip().lower() in ["pass", "passed", "true", "yes", "y", "1"]
        else:
            passed = bool(res_val)
        status_text = "PASSED" if passed else "FAILED"
        details = res.get("details", "")
        # Convert details dict to string if needed
        if isinstance(details, dict):
            # Format each key: value; for readability
            detail_parts = []
            for k, v in details.items():
                # Format values (truncate floats, represent booleans clearly)
                if isinstance(v, float):
                    v_str = f"{v:.3f}"  # format float to 3 decimal places
                elif isinstance(v, bool):
                    v_str = "True" if v else "False"
                else:
                    v_str = str(v)
                detail_parts.append(f"{k}: {v_str}")
            details_text = "; ".join(detail_parts)
        else:
            details_text = str(details)
        # Add a new row to the table
        row_cells = table.add_row().cells
        row_cells[0].text = str(test_id)
        row_cells[1].text = str(req_id)
        # For result, use bold text and color for emphasis
        result_run = row_cells[2].paragraphs[0].add_run(status_text)
        result_run.bold = True
        if passed:
            # Optionally color passed as green (not mandatory)
            # result_run.font.color.rgb = RGBColor(0, 128, 0)
            pass
        else:
            # Color failed results in red for high visibility
            result_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        # Details cell
        row_cells[3].text = details_text
    # Save the document
    try:
        document.save(output_path)
    except Exception as e:
        logging.error(f"Failed to save the report: {e}")
        sys.exit(1)
    logging.info(f"Compliance report generated successfully: {output_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate a compliance report from test results.")
    parser.add_argument("--input", required=True, help="Path to test results (directory of JSONs, or a JSON/CSV file).")
    parser.add_argument("--output", required=False, help="Output report file name (should end with .docx).")
    parser.add_argument("--template", required=False, help="Path to a Word template .docx (optional).")
    args = parser.parse_args()

    input_path = args.input
    output_path = args.output or "compliance_report.docx"
    template_path = args.template

    # If output given ends with .pdf, inform user that we'll create a .docx instead (no direct PDF support)
    if output_path.lower().endswith(".pdf"):
        logging.warning("Output file is a .pdf; the report will be generated as .docx and should be converted to PDF separately.")
        output_path = output_path[:-4] + ".docx"
    elif not output_path.lower().endswith(".docx"):
        output_path += ".docx"

    # Load test results
    results = load_test_results(input_path)
    if not results:
        logging.error("No test results found to report on. Exiting.")
        sys.exit(1)
    # Generate the compliance report document
    generate_report(results, output_path, template_path)
    logging.info("Compliance report generation script completed.")
